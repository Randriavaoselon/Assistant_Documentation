"""
Extraction du texte brut à partir de différents formats de fichiers
(PDF, DOCX, TXT, Markdown).
"""
from pathlib import Path


class ExtractionError(Exception):
    pass


def extract_text(file_path: str, file_type: str) -> str:
    """Retourne le texte brut extrait du fichier, selon son type."""
    file_type = file_type.lower().lstrip(".")

    if file_type == "pdf":
        return _extract_pdf(file_path)
    if file_type == "docx":
        return _extract_docx(file_path)
    if file_type in ("txt", "md", "markdown"):
        return _extract_plain_text(file_path)

    raise ExtractionError(
        f"Type de fichier non supporté : '{file_type}'. "
        "Formats acceptés : PDF, DOCX, TXT, MD."
    )


def _extract_pdf(file_path: str) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(file_path)
    except Exception as exc:
        raise ExtractionError(f"Impossible de lire le PDF : {exc}") from exc

    pages_text = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(f"[Page {page_number}]\n{text.strip()}")

    full_text = "\n\n".join(pages_text)
    if not full_text.strip():
        raise ExtractionError(
            "Aucun texte n'a pu être extrait de ce PDF. "
            "S'agit-il d'un PDF scanné sans OCR ?"
        )
    return full_text


def _extract_docx(file_path: str) -> str:
    import docx

    try:
        doc = docx.Document(file_path)
    except Exception as exc:
        raise ExtractionError(f"Impossible de lire le DOCX : {exc}") from exc

    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)
    if not full_text.strip():
        raise ExtractionError("Aucun texte n'a pu être extrait de ce DOCX.")
    return full_text


def _extract_plain_text(file_path: str) -> str:
    path = Path(file_path)
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")
